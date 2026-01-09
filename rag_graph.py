import os
import zipfile
import json
import xml.etree.ElementTree as ET
from typing import Annotated, List, TypedDict, Union
from dotenv import load_dotenv
import pandas as pd

# LangChain imports
from langchain_openai import ChatOpenAI
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_core.messages import BaseMessage, HumanMessage, AIMessage
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.runnables import RunnableConfig

# LangGraph imports
from langgraph.graph import StateGraph, END

# Load environment variables from .env file
load_dotenv()

# Check for API Key
if not os.getenv("DEEPSEEK_API_KEY"):
    print("Warning: DEEPSEEK_API_KEY not found in environment variables. Please set it in .env file.")

# --- 1. Setup Vector Store (RAG) ---
def load_documents_from_directory(directory_path):
    documents = []
    if not os.path.exists(directory_path):
        print(f"Directory not found: {directory_path}")
        return documents
        
    for filename in os.listdir(directory_path):
        # Skip temporary files
        if filename.startswith("~$"):
            continue
            
        file_path = os.path.join(directory_path, filename)
        if filename.endswith(('.xlsx', '.xls')):
            try:
                df = pd.read_excel(file_path)
                # Convert each row to a document
                for index, row in df.iterrows():
                    # Combine all columns into text content
                    content = "\n".join([f"{col}: {val}" for col, val in row.items() if pd.notna(val)])
                    documents.append(Document(page_content=content, metadata={"source": filename, "row": index}))
                print(f"Loaded {len(df)} rows from {filename}")
            except Exception as e:
                print(f"Error loading Excel file {filename}: {e}")
        elif filename.endswith('.txt'):
            content = None
            # Try different encodings
            encodings = ['utf-8', 'gb18030', 'gbk', 'utf-16']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    print(f"Loaded text file {filename} (using {encoding})")
                    break
                except UnicodeDecodeError:
                    continue
                except Exception as e:
                    print(f"Error loading text file {filename} with {encoding}: {e}")
                    break
            
            # If all strict encodings fail, try with error ignoring
            if content is None:
                try:
                    print(f"Warning: Could not decode {filename} with standard encodings. Trying to ignore errors...")
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        content = f.read()
                    print(f"Loaded text file {filename} (with errors ignored)")
                except Exception as e:
                    print(f"Failed to load text file {filename}: {e}")

            if content:
                documents.append(Document(page_content=content, metadata={"source": filename}))
        elif filename.endswith('.docx'):
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument(file_path)
                content = "\n".join([para.text for para in doc.paragraphs])
                documents.append(Document(page_content=content, metadata={"source": filename}))
                print(f"Loaded Word document {filename}")
            except Exception as e:
                print(f"Standard load failed for {filename}, trying fallback...")
                try:
                    # Fallback: Try to read word/document.xml directly from zip
                    with zipfile.ZipFile(file_path) as z:
                        xml_content = z.read('word/document.xml')
                        tree = ET.fromstring(xml_content)
                        # Simple text extraction from XML
                        text_parts = []
                        for node in tree.iter():
                            if node.tag.endswith('}t'): # Text node
                                if node.text:
                                    text_parts.append(node.text)
                            elif node.tag.endswith('}p'): # Paragraph end
                                text_parts.append('\n')
                        content = "".join(text_parts).strip()
                        if content:
                            documents.append(Document(page_content=content, metadata={"source": filename}))
                            print(f"Loaded Word document {filename} (fallback)")
                        else:
                            print(f"Fallback loaded empty content for {filename}")
                except Exception as e2:
                    print(f"Error loading Word file {filename}: {e} | Fallback error: {e2}")
        elif filename.endswith('.json'):
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                
                # Handle PsyQA format (list of dicts)
                if isinstance(data, list):
                    count = 0
                    for item in data:
                        # Extract basic fields
                        question = item.get("question", "")
                        description = item.get("description", "")
                        keywords = item.get("keywords", "")
                        
                        # Extract answers
                        answers = item.get("answers", [])
                        answer_texts = []
                        for ans in answers:
                            if isinstance(ans, dict) and "answer_text" in ans:
                                answer_texts.append(ans["answer_text"])
                        
                        # Construct content block
                        # Combining Question + Description + Answers into one context block
                        content_parts = []
                        if question: content_parts.append(f"问题: {question}")
                        if description: content_parts.append(f"详细描述: {description}")
                        if keywords: content_parts.append(f"关键词: {keywords}")
                        if answer_texts: content_parts.append(f"专家回答: {' '.join(answer_texts)}")
                        
                        content = "\n".join(content_parts)
                        
                        if content:
                            documents.append(Document(page_content=content, metadata={"source": filename}))
                            count += 1
                    print(f"Loaded {count} QA pairs from JSON file {filename}")
                else:
                    print(f"JSON structure not recognized for {filename} (expected a list)")
            except Exception as e:
                print(f"Error loading JSON file {filename}: {e}")
    return documents

print("Initializing Vector Store...")
# Initialize Embeddings
# Switched to a better Chinese embedding model
try:
    embeddings = HuggingFaceEmbeddings(model_name="shibing624/text2vec-base-chinese")
    vector_store_path = "faiss_index"

    if os.path.exists(vector_store_path):
        print(f"Loading existing vector store from {vector_store_path}...")
        vector_store = FAISS.load_local(vector_store_path, embeddings, allow_dangerous_deserialization=True)
    else:
        print("Index not found. Loading documents and creating new vector store...")
        # Load documents from data directories
        knowledge_base_docs = load_documents_from_directory(os.path.join("data", "knowledge_base"))
        cases_docs = load_documents_from_directory(os.path.join("data", "cases"))
        raw_documents = knowledge_base_docs + cases_docs

        if not raw_documents:
            print("No documents found in data directories. Using dummy data for testing.")
            raw_documents = [
                Document(page_content="LangGraph is a library for building stateful, multi-actor applications with LLMs, built on top of LangChain."),
                Document(page_content="LangChain is a framework for developing applications powered by language models. It enables applications that are context-aware and reason."),
                Document(page_content="Retrieval-Augmented Generation (RAG) is a technique for enhancing the accuracy and reliability of generative AI models with facts fetched from external sources."),
                Document(page_content="The user is asking for a python script using langchain and langgraph."),
            ]
        
        # Split documents into chunks for better retrieval
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=500,
            chunk_overlap=50,
            separators=["\n\n", "\n", "。", "！", "？", " ", ""]
        )
        all_documents = text_splitter.split_documents(raw_documents)
        print(f"Splitting documents: {len(raw_documents)} raw docs -> {len(all_documents)} chunks")
        
        vector_store = FAISS.from_documents(all_documents, embeddings)
        vector_store.save_local(vector_store_path)
        print(f"Vector store created and saved to {vector_store_path}")

    retriever = vector_store.as_retriever(search_kwargs={"k": 3})
except Exception as e:
    print(f"Error initializing vector store: {e}")
    vector_store = None
    retriever = None

# --- 2. Define Graph State ---
class AgentState(TypedDict):
    messages: List[BaseMessage]
    context: str

# --- 3. Define Nodes ---

def retrieve_node(state: AgentState):
    """
    Retrieve relevant documents based on the last user message.
    """
    print("--- Node: Retrieve ---")
    messages = state['messages']
    last_message = messages[-1]
    query = last_message.content
    
    if retriever:
        docs = retriever.invoke(query)
        context = "\n\n".join([doc.page_content for doc in docs])
    else:
        context = "No context available (Vector store not initialized)."
    
    print(f"Retrieved context length: {len(context)}")
    # print(f"Retrieved context content:\n{context}\n--------------------")
    return {"context": context}

def generate_node(state: AgentState):
    """
    Generate a response using the LLM and the retrieved context.
    """
    print("--- Node: Generate ---")
    messages = state['messages']
    context = state['context']
    
    # Read system prompt from file
    system_prompt_path = os.path.join("prompts", "system_prompt.txt")
    default_prompt = "You are a helpful assistant. Use the following context to answer the user's question.\n\nContext:\n{context}"
    
    if os.path.exists(system_prompt_path):
        try:
            with open(system_prompt_path, "r", encoding="utf-8") as f:
                system_prompt_text = f.read()
            # Ensure {context} placeholder exists so RAG works
            if "{context}" not in system_prompt_text:
                system_prompt_text += "\n\n相关上下文:\n{context}"
        except Exception as e:
            print(f"Error reading system prompt file: {e}")
            system_prompt_text = default_prompt
    else:
        system_prompt_text = default_prompt

    # Define the prompt template
    prompt = ChatPromptTemplate.from_messages([
        ("system", system_prompt_text),
        ("placeholder", "{messages}"),
    ])
    
    # Initialize LLM
    # Using DeepSeek API via ChatOpenAI client
    model = ChatOpenAI(
        model="deepseek-chat", 
        temperature=0,
        api_key=os.getenv("DEEPSEEK_API_KEY"),
        base_url=os.getenv("DEEPSEEK_BASE_URL")
    )
    
    # Create chain
    chain = prompt | model
    
    # Invoke chain
    response = chain.invoke({"context": context, "messages": messages})
    
    return {"messages": [response]}

# --- 4. Build Graph ---
workflow = StateGraph(AgentState)

# Add nodes
workflow.add_node("retrieve", retrieve_node)
workflow.add_node("generate", generate_node)

# Define edges
workflow.set_entry_point("retrieve")
workflow.add_edge("retrieve", "generate")
workflow.add_edge("generate", END)

# Compile the graph
app = workflow.compile()

# --- 5. Execution (Main) ---
if __name__ == "__main__":
    print("\n--- 心理咨询助手已启动 (输入 'quit' 或 'exit' 退出) ---\n")
    
    while True:
        user_query = input("User: ")
        if user_query.lower() in ["quit", "exit"]:
            print("Goodbye!")
            break
            
        if not user_query.strip():
            continue

        # Initial State
        initial_state = {
            "messages": [HumanMessage(content=user_query)],
            "context": ""
        }
        
        try:
            # Stream the graph execution
            for output in app.stream(initial_state):
                for key, value in output.items():
                    if key == "generate":
                        last_msg = value["messages"][-1]
                        print(f"\nAssistant: {last_msg.content}\n")
        except Exception as e:
            print(f"An error occurred: {e}")
            print("Ensure you have set DEEPSEEK_API_KEY in .env and installed requirements.")
